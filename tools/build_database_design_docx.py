"""生成数据库设计 Word 文档，并统一控制页面、表格、字体和页眉页脚格式。"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "database-design"
OUTPUT_PATH = OUTPUT_DIR / "个人金融NL2SQL营销查询系统-数据库设计.docx"
ARCH_IMAGE = Path(r"C:\Users\30988\AppData\Local\Temp\codex-clipboard-268bf9fa-512a-4260-802e-23312f396b6f.png")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
VERY_LIGHT = "F4F6F9"
PALE_GOLD = "FFF8E8"
WHITE = "FFFFFF"
GRID = "AEB8C4"
TABLE_FONT_SIZE = 8.3


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=None, color=None, name="Calibri", east_asia="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph_runs(paragraph, size=11, color=None, bold=None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def add_bottom_border(paragraph, color="D7DBE2", size="6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_section_header_footer(section) -> None:
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run("个人金融 NL2SQL 营销查询系统")
    set_run_font(r1, size=8.5, bold=True, color=INK)
    r2 = p.add_run("  |  数据库设计")
    set_run_font(r2, size=8.5, color=MUTED)
    add_bottom_border(p)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("内部设计文档  |  第 ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)
    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=9, color=MUTED)


def configure_section(section, landscape=False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85 if landscape else 1.0)
    section.right_margin = Inches(0.85 if landscape else 1.0)
    set_section_header_footer(section)


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract_ids or [0]) + 1
    num_id = max(existing_num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    r = p.add_run(text)
    set_run_font(r, size=11)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_callout(doc, label, text, fill=VERY_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="D7DBE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, size=10.5, bold=True, color=INK)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=TABLE_FONT_SIZE,
              center_cols=None, landscape=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, (cell, text) in enumerate(zip(hdr.cells, headers)):
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(str(text))
        set_run_font(r, size=font_size, bold=True, color=INK)
    center_cols = set(center_cols or [])
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, (cell, value) in enumerate(zip(row.cells, row_values)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in center_cols else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run("" if value is None else str(value))
            set_run_font(r, size=font_size)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)
    return p


def add_code_block(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="CBD3DC", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for i, line in enumerate(code.strip().splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line.rstrip())
        set_run_font(r, size=8.2, name="Consolas", east_asia="Microsoft YaHei", color="2B2F33")
    return table


def set_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


TABLES = [
    {
        "cn": "客户表",
        "name": "dim_customer",
        "desc": "客户统一画像维表。保存个人客户基础属性、身份等级、风险与 KYC 状态、所属机构及当前客户经理关系，为查客户、客户分层、营销客群筛选提供统一客户口径。",
        "grain": "每个 customer_id、每个有效版本一行；snapshot_dt 分区内保留当日有效快照。",
        "refresh": "每日全量快照；客户等级、经理关系等字段按 SCD2 维护有效期。",
        "key": "逻辑主键 customer_sk；业务唯一键 customer_id + effective_start_time；分区键 snapshot_dt。",
        "partition": "PARTITIONED BY (snapshot_dt DATE)；建议按日保留最近 13 个月热分区，历史转冷存储。",
        "accel": "按 customer_id 分桶 128 桶；桶内按 customer_id、is_current 排序；ORC Bloom Filter 建议覆盖 customer_id、manager_id；每日收集列统计。",
        "fields": [
            ("customer_sk", "BIGINT", "是", "PK", "客户代理键，仅用于数仓内部稳定关联", "10000001"),
            ("customer_id", "STRING", "是", "UK/桶键", "脱敏后的全局客户标识；跨表关联统一使用", "C9F3..."),
            ("customer_no", "STRING", "否", "查询键", "客户号密文或令牌，不保存明文核心客户号", "TKN_8A..."),
            ("customer_name_masked", "STRING", "否", "-", "脱敏客户姓名，仅用于授权展示", "张*"),
            ("gender_code", "STRING", "否", "-", "性别代码：M/F/U；以码表为准", "F"),
            ("birth_date", "DATE", "否", "敏感列", "出生日期；仅在授权范围内可见", "1988-05-12"),
            ("age", "SMALLINT", "否", "派生", "按 snapshot_dt 计算的周岁，不作为永久静态属性", "38"),
            ("age_band_code", "STRING", "否", "查询键", "年龄段代码，如 A18_25、A26_35", "A36_45"),
            ("mobile_masked", "STRING", "否", "敏感列", "脱敏手机号；禁止保存完整号码", "138****5678"),
            ("id_type_code", "STRING", "否", "-", "证件类型代码", "ID_CARD"),
            ("id_no_hash", "STRING", "否", "查询键", "证件号码加盐哈希，用于去重，不可逆", "SHA256..."),
            ("id_no_last4", "STRING", "否", "敏感列", "证件号码后四位，授权展示使用", "1234"),
            ("customer_level_code", "STRING", "是", "查询键", "客户等级代码，如 NORMAL/GOLD/PLATINUM", "GOLD"),
            ("customer_type_code", "STRING", "是", "-", "客户类型；本系统固定为 PERSONAL，预留扩展", "PERSONAL"),
            ("vip_flag", "BOOLEAN", "是", "查询键", "是否 VIP 客户", "true"),
            ("risk_level_code", "STRING", "否", "查询键", "客户风险承受等级，如 R1-R5", "R3"),
            ("kyc_status_code", "STRING", "是", "查询键", "KYC 状态：VALID/EXPIRING/EXPIRED/UNKNOWN", "VALID"),
            ("occupation_code", "STRING", "否", "查询键", "职业代码，来自统一码表", "FINANCE"),
            ("annual_income_band_code", "STRING", "否", "查询键", "年收入区间，不落精确收入", "I30_50W"),
            ("asset_level_code", "STRING", "否", "查询键", "金融资产等级，口径由资产分层规则定义", "A50_100W"),
            ("branch_id", "STRING", "是", "查询键", "客户归属网点/机构标识", "B00102"),
            ("manager_id", "STRING", "否", "查询键", "当前客户经理标识，关联 dim_customer_manager", "M000128"),
            ("open_date", "DATE", "否", "-", "客户首次建立关系日期", "2019-03-08"),
            ("status_code", "STRING", "是", "查询键", "客户状态：ACTIVE/INACTIVE/CLOSED", "ACTIVE"),
            ("preferred_contact_code", "STRING", "否", "-", "首选联系渠道：PHONE/SMS/APP/WECHAT/NONE", "APP"),
            ("province_code", "STRING", "否", "查询键", "常驻省级行政区代码", "440000"),
            ("city_code", "STRING", "否", "查询键", "常驻地市代码", "440300"),
            ("effective_start_time", "TIMESTAMP", "是", "SCD2", "该版本生效时间", "2026-08-01 00:00:00"),
            ("effective_end_time", "TIMESTAMP", "是", "SCD2", "该版本失效时间；当前版本取 9999-12-31", "9999-12-31 23:59:59"),
            ("is_current", "BOOLEAN", "是", "查询键", "是否当前有效版本", "true"),
            ("source_system_code", "STRING", "是", "审计", "来源系统代码", "CRM"),
            ("etl_batch_id", "STRING", "是", "审计", "ETL 批次号，支持血缘与回溯", "20260825_01"),
            ("data_update_time", "TIMESTAMP", "是", "审计", "记录在数仓中的最后更新时间", "2026-08-25 02:10:00"),
            ("snapshot_dt", "DATE", "是", "分区键", "快照业务日期，格式 yyyy-MM-dd", "2026-08-25"),
        ],
    },
    {
        "cn": "交易主题表",
        "name": "fct_transaction",
        "desc": "个人客户交易明细主题事实表。承载账户交易的日期、金额、类型、渠道和余额等信息，支持交易查询、统计、客户活跃度和营销触发分析。",
        "grain": "每一笔入账后的交易流水一行；冲正交易单独记录并通过 original_transaction_id 关联。",
        "refresh": "准实时/小时增量入湖，T+1 对账补数；按 txn_dt 幂等重跑。",
        "key": "逻辑主键 transaction_id；分区键 txn_dt；customer_id 为主要关联键。",
        "partition": "PARTITIONED BY (txn_dt DATE)；按交易日期裁剪，避免对 transaction_time 直接函数过滤。",
        "accel": "按 customer_id 分桶 256 桶；桶内按 customer_id、transaction_time 排序；对 transaction_id、customer_id 启用 Bloom Filter；高频统计使用汇总视图。",
        "fields": [
            ("transaction_id", "STRING", "是", "PK", "全局唯一交易流水标识；重跑去重依据", "TXN2026..."),
            ("customer_id", "STRING", "是", "FK/桶键", "客户标识，关联 dim_customer", "C9F3..."),
            ("account_id_hash", "STRING", "是", "查询键", "账户标识加盐哈希，不保存明文账号", "SHA256..."),
            ("product_id", "STRING", "否", "查询键", "交易关联产品标识；产品主数据由外部维表提供", "P_DEP_001"),
            ("transaction_time", "TIMESTAMP", "是", "排序键", "交易发生时间，统一存储为 Asia/Hong_Kong 业务时间", "2026-08-25 10:32:08"),
            ("transaction_date", "DATE", "是", "查询键", "交易业务日期，与分区日期原则上一致", "2026-08-25"),
            ("transaction_type_code", "STRING", "是", "查询键", "交易大类：TRANSFER/CONSUME/DEPOSIT/WITHDRAW/INTEREST 等", "CONSUME"),
            ("transaction_subtype_code", "STRING", "否", "查询键", "交易子类代码，来自统一交易码表", "POS_PURCHASE"),
            ("channel_code", "STRING", "否", "查询键", "渠道：APP/WEB/ATM/POS/COUNTER/AUTO", "APP"),
            ("debit_credit_flag", "STRING", "是", "查询键", "借贷方向：D 支出，C 收入", "D"),
            ("currency_code", "STRING", "是", "查询键", "ISO 4217 币种代码", "CNY"),
            ("amount", "DECIMAL(20,2)", "是", "度量", "原币交易金额，非负；方向由 debit_credit_flag 表示", "1288.00"),
            ("amount_cny", "DECIMAL(20,2)", "是", "度量", "按交易日标准汇率折算的人民币金额", "1288.00"),
            ("balance_after", "DECIMAL(20,2)", "否", "度量", "交易后账户余额；无余额来源时置空", "52100.36"),
            ("counterparty_type_code", "STRING", "否", "查询键", "对手方类型：PERSON/COMPANY/MERCHANT/BANK/UNKNOWN", "MERCHANT"),
            ("counterparty_name_masked", "STRING", "否", "敏感列", "脱敏交易对手名称", "深***公司"),
            ("merchant_category_code", "STRING", "否", "查询键", "商户类别码 MCC", "5411"),
            ("branch_id", "STRING", "否", "查询键", "受理机构/网点标识", "B00102"),
            ("status_code", "STRING", "是", "查询键", "交易状态：SUCCESS/FAILED/PENDING/REVERSED", "SUCCESS"),
            ("reversal_flag", "BOOLEAN", "是", "查询键", "是否冲正记录", "false"),
            ("original_transaction_id", "STRING", "否", "关联键", "冲正对应的原交易流水号", "TXN2026..."),
            ("summary_masked", "STRING", "否", "敏感列", "交易摘要脱敏文本；不得包含完整账号、手机号", "手机银行消费"),
            ("source_system_code", "STRING", "是", "审计", "来源系统代码", "CORE_BANK"),
            ("etl_batch_id", "STRING", "是", "审计", "ETL 批次号", "20260825_10"),
            ("ingestion_time", "TIMESTAMP", "是", "审计", "进入数仓的时间", "2026-08-25 10:35:00"),
            ("txn_dt", "DATE", "是", "分区键", "交易分区日期", "2026-08-25"),
        ],
    },
    {
        "cn": "营销活动表",
        "name": "dim_marketing_campaign",
        "desc": "营销活动定义维表。描述活动名称、目标客群、渠道、关联产品、预算、优惠和有效期，为营销查询和效果评估提供统一活动口径。",
        "grain": "每个 campaign_id、每个有效版本一行；活动定义变更按 SCD2 留痕。",
        "refresh": "活动变更触发增量同步，每日补偿全量；按 snapshot_dt 保留当日快照。",
        "key": "逻辑主键 campaign_sk；业务唯一键 campaign_id + effective_start_time。",
        "partition": "PARTITIONED BY (snapshot_dt DATE)；活动数量较小时不建议过度分桶。",
        "accel": "按 campaign_id 分桶 32 桶；按 campaign_status_code、start_time 排序；对 campaign_id、product_id 收集列统计。",
        "fields": [
            ("campaign_sk", "BIGINT", "是", "PK", "活动代理键", "300001"),
            ("campaign_id", "STRING", "是", "UK/桶键", "营销活动业务标识", "CMP202608001"),
            ("campaign_name", "STRING", "是", "查询键", "营销活动名称；应避免包含敏感客户信息", "财富季基金定投活动"),
            ("campaign_type_code", "STRING", "是", "查询键", "活动类型：ACQUISITION/UPSELL/RETENTION/CARE", "UPSELL"),
            ("campaign_status_code", "STRING", "是", "查询键", "DRAFT/APPROVED/RUNNING/PAUSED/FINISHED/CANCELLED", "RUNNING"),
            ("product_id", "STRING", "否", "关联键", "主推产品标识；多产品活动可由扩展关系表承载", "P_FUND_023"),
            ("target_customer_segment_code", "STRING", "否", "查询键", "目标客群分层代码", "AUM_50W_PLUS"),
            ("channel_code", "STRING", "是", "查询键", "主触达渠道；多渠道可使用 MULTI", "APP"),
            ("owner_org_id", "STRING", "是", "查询键", "活动主责机构标识", "ORG001"),
            ("owner_manager_id", "STRING", "否", "关联键", "活动负责人客户经理/员工标识", "M000128"),
            ("start_time", "TIMESTAMP", "是", "查询键", "活动开始时间", "2026-08-01 00:00:00"),
            ("end_time", "TIMESTAMP", "是", "查询键", "活动结束时间", "2026-09-30 23:59:59"),
            ("budget_amount", "DECIMAL(20,2)", "否", "度量", "活动预算金额，人民币", "500000.00"),
            ("target_count", "BIGINT", "否", "度量", "计划覆盖客户数", "50000"),
            ("offer_type_code", "STRING", "否", "查询键", "优惠类型：RATE/COUPON/GIFT/FEE_REDUCTION/NONE", "COUPON"),
            ("offer_value", "DECIMAL(20,4)", "否", "度量", "优惠值；由 offer_type_code 决定单位", "50.0000"),
            ("contact_frequency_limit", "INT", "是", "规则", "活动周期内单客户最大触达次数", "3"),
            ("consent_required_flag", "BOOLEAN", "是", "规则", "触达前是否必须具备营销授权", "true"),
            ("objective_desc", "STRING", "否", "-", "活动目标简述，限制长度并做内容审核", "提升基金定投参与率"),
            ("effective_start_time", "TIMESTAMP", "是", "SCD2", "版本生效时间", "2026-07-20 09:00:00"),
            ("effective_end_time", "TIMESTAMP", "是", "SCD2", "版本失效时间", "9999-12-31 23:59:59"),
            ("is_current", "BOOLEAN", "是", "查询键", "是否当前有效版本", "true"),
            ("source_system_code", "STRING", "是", "审计", "来源系统代码", "MARKETING"),
            ("etl_batch_id", "STRING", "是", "审计", "ETL 批次号", "20260825_02"),
            ("data_update_time", "TIMESTAMP", "是", "审计", "数仓更新时间", "2026-08-25 02:30:00"),
            ("snapshot_dt", "DATE", "是", "分区键", "快照业务日期", "2026-08-25"),
        ],
    },
    {
        "cn": "产品持有表",
        "name": "fct_product_holding",
        "desc": "客户产品持有快照事实表。记录存款、理财、基金等产品的账户级持仓、余额、市值、收益、到期日和风险等级，支撑持仓查询和客户资产分析。",
        "grain": "每个 customer_id + account_id_hash + product_id 在每个 snapshot_dt 一行。",
        "refresh": "每日 T+1 全量快照；对盘中场景可增加小时级临时表，但不改变本表日终口径。",
        "key": "业务唯一键 customer_id + account_id_hash + product_id + snapshot_dt。",
        "partition": "PARTITIONED BY (snapshot_dt DATE)；查询必须优先命中最新或指定快照分区。",
        "accel": "按 customer_id 分桶 256 桶；桶内按 customer_id、product_category_code、product_id 排序；对 customer_id、product_id 启用 Bloom Filter。",
        "fields": [
            ("holding_id", "STRING", "是", "PK", "持仓记录稳定标识，可由业务键哈希生成", "HLD..."),
            ("customer_id", "STRING", "是", "FK/桶键", "客户标识", "C9F3..."),
            ("account_id_hash", "STRING", "是", "UK", "账户标识加盐哈希", "SHA256..."),
            ("product_id", "STRING", "是", "UK/查询键", "产品标识，关联外部产品主数据", "P_WM_010"),
            ("product_category_code", "STRING", "是", "查询键", "DEPOSIT/WEALTH/FUND/INSURANCE/BOND/OTHER", "WEALTH"),
            ("holding_status_code", "STRING", "是", "查询键", "ACTIVE/MATURED/REDEEMED/FROZEN", "ACTIVE"),
            ("open_date", "DATE", "否", "查询键", "开户或首次持有日期", "2025-11-20"),
            ("maturity_date", "DATE", "否", "查询键", "到期日期；无固定期限时置空", "2026-11-20"),
            ("currency_code", "STRING", "是", "查询键", "ISO 4217 币种代码", "CNY"),
            ("principal_amount", "DECIMAL(20,2)", "否", "度量", "本金或累计申购金额", "100000.00"),
            ("balance_amount", "DECIMAL(20,2)", "否", "度量", "账面余额", "102350.18"),
            ("market_value", "DECIMAL(20,2)", "否", "度量", "估值市值；存款可等于余额", "103200.00"),
            ("available_balance", "DECIMAL(20,2)", "否", "度量", "可用余额或可赎回金额", "100000.00"),
            ("accumulated_income", "DECIMAL(20,2)", "否", "度量", "累计收益，允许为负", "3200.00"),
            ("interest_rate", "DECIMAL(9,6)", "否", "度量", "年化利率/收益率小数，例如 0.032500", "0.032500"),
            ("term_days", "INT", "否", "查询键", "产品期限天数；开放式产品置空", "365"),
            ("risk_level_code", "STRING", "否", "查询键", "产品风险等级，如 PR1-PR5", "PR3"),
            ("auto_renew_flag", "BOOLEAN", "否", "查询键", "是否自动续期", "false"),
            ("branch_id", "STRING", "否", "查询键", "销售/归属机构标识", "B00102"),
            ("manager_id", "STRING", "否", "查询键", "服务客户经理标识", "M000128"),
            ("source_system_code", "STRING", "是", "审计", "来源系统代码", "ASSET"),
            ("etl_batch_id", "STRING", "是", "审计", "ETL 批次号", "20260825_EOD"),
            ("data_update_time", "TIMESTAMP", "是", "审计", "数仓更新时间", "2026-08-25 04:10:00"),
            ("snapshot_dt", "DATE", "是", "分区键", "日终快照业务日期", "2026-08-25"),
        ],
    },
    {
        "cn": "客户经理表",
        "name": "dim_customer_manager",
        "desc": "客户经理维表。记录客户经理的脱敏身份、所属机构、岗位、等级、资质、服务状态和管理规模，用于数据权限、客户归属和绩效查询。",
        "grain": "每个 manager_id、每个有效版本一行；组织或岗位变更按 SCD2 留痕。",
        "refresh": "人力/CRM 变更增量同步，每日全量核对；snapshot_dt 保存有效快照。",
        "key": "逻辑主键 manager_sk；业务唯一键 manager_id + effective_start_time。",
        "partition": "PARTITIONED BY (snapshot_dt DATE)；规模较小时按日分区但不做高桶数。",
        "accel": "按 manager_id 分桶 32 桶；按 org_id、service_status_code 排序；对 manager_id、org_id 收集列统计。",
        "fields": [
            ("manager_sk", "BIGINT", "是", "PK", "客户经理代理键", "500001"),
            ("manager_id", "STRING", "是", "UK/桶键", "客户经理统一标识", "M000128"),
            ("employee_no_hash", "STRING", "否", "查询键", "员工号加盐哈希", "SHA256..."),
            ("manager_name_masked", "STRING", "是", "敏感列", "脱敏姓名，用于授权展示", "李*"),
            ("mobile_masked", "STRING", "否", "敏感列", "脱敏办公手机号", "139****1024"),
            ("org_id", "STRING", "是", "查询键", "所属组织机构标识", "ORG001"),
            ("branch_id", "STRING", "是", "查询键", "所属网点标识", "B00102"),
            ("post_code", "STRING", "是", "查询键", "岗位代码，如 RM/WEALTH_RM/TEAM_LEAD", "WEALTH_RM"),
            ("rank_code", "STRING", "否", "查询键", "客户经理等级代码", "SENIOR"),
            ("certificate_level_code", "STRING", "否", "查询键", "专业资质等级或组合代码", "FUND_INSURANCE"),
            ("service_status_code", "STRING", "是", "查询键", "ACTIVE/LEAVE/TRANSFERRED/INACTIVE", "ACTIVE"),
            ("region_code", "STRING", "否", "查询键", "服务区域代码", "SZ_NANSHAN"),
            ("customer_count", "BIGINT", "否", "度量", "snapshot_dt 当日名下有效客户数", "326"),
            ("asset_aum_amount", "DECIMAL(20,2)", "否", "度量", "snapshot_dt 当日名下客户资产规模，人民币", "128500000.00"),
            ("start_date", "DATE", "否", "-", "上岗日期", "2021-06-01"),
            ("end_date", "DATE", "否", "-", "离岗日期；在岗时为空", ""),
            ("effective_start_time", "TIMESTAMP", "是", "SCD2", "版本生效时间", "2026-07-01 00:00:00"),
            ("effective_end_time", "TIMESTAMP", "是", "SCD2", "版本失效时间", "9999-12-31 23:59:59"),
            ("is_current", "BOOLEAN", "是", "查询键", "是否当前有效版本", "true"),
            ("source_system_code", "STRING", "是", "审计", "来源系统代码", "HR_CRM"),
            ("etl_batch_id", "STRING", "是", "审计", "ETL 批次号", "20260825_01"),
            ("data_update_time", "TIMESTAMP", "是", "审计", "数仓更新时间", "2026-08-25 03:00:00"),
            ("snapshot_dt", "DATE", "是", "分区键", "快照业务日期", "2026-08-25"),
        ],
    },
    {
        "cn": "客户营销关联表",
        "name": "fct_customer_marketing",
        "desc": "客户与营销活动的触达、参与、响应和转化事件事实表。每次营销事件独立记录，用于活动漏斗、客户跟进、响应分析和转化归因。",
        "grain": "每个客户在某活动下的一次原子营销事件一行；同一客户可对应多次触达与响应。",
        "refresh": "事件流分钟级写入，T+1 与渠道回执对账；按 event_dt 幂等去重。",
        "key": "逻辑主键 interaction_id；客户、活动、事件时间构成主要查询路径。",
        "partition": "PARTITIONED BY (event_dt DATE)；跨期活动查询必须显式限定日期范围。",
        "accel": "按 campaign_id 分桶 128 桶，桶内按 campaign_id、customer_id、event_time 排序；对 interaction_id、customer_id、campaign_id 启用 Bloom Filter。",
        "fields": [
            ("interaction_id", "STRING", "是", "PK", "营销事件唯一标识；事件去重依据", "INT2026..."),
            ("customer_id", "STRING", "是", "FK/查询键", "客户标识", "C9F3..."),
            ("campaign_id", "STRING", "是", "FK/桶键", "活动标识", "CMP202608001"),
            ("manager_id", "STRING", "否", "FK/查询键", "执行触达或跟进的客户经理标识", "M000128"),
            ("product_id", "STRING", "否", "查询键", "本次事件关联产品标识", "P_FUND_023"),
            ("event_time", "TIMESTAMP", "是", "排序键", "事件发生时间", "2026-08-25 11:00:00"),
            ("event_date", "DATE", "是", "查询键", "事件业务日期", "2026-08-25"),
            ("event_type_code", "STRING", "是", "查询键", "TARGET/EXPOSURE/CONTACT/CLICK/REGISTER/PURCHASE/FOLLOW_UP", "CONTACT"),
            ("channel_code", "STRING", "是", "查询键", "触达或响应渠道", "PHONE"),
            ("contact_batch_no", "STRING", "否", "查询键", "批量触达批次号", "BATCH0825A"),
            ("contact_result_code", "STRING", "否", "查询键", "SUCCESS/NO_ANSWER/REJECT/INVALID/FAILED", "SUCCESS"),
            ("response_flag", "BOOLEAN", "是", "查询键", "客户是否产生有效响应", "true"),
            ("response_time", "TIMESTAMP", "否", "-", "首次有效响应时间", "2026-08-25 11:05:00"),
            ("response_type_code", "STRING", "否", "查询键", "INTERESTED/NOT_INTERESTED/NEED_FOLLOW_UP/COMPLAINT", "INTERESTED"),
            ("conversion_flag", "BOOLEAN", "是", "查询键", "是否满足活动定义的转化规则", "false"),
            ("conversion_time", "TIMESTAMP", "否", "-", "转化发生时间", ""),
            ("conversion_amount", "DECIMAL(20,2)", "否", "度量", "归因到活动的转化金额，人民币", "0.00"),
            ("consent_status_code", "STRING", "是", "合规", "事件发生时的营销授权状态：GRANTED/DENIED/UNKNOWN", "GRANTED"),
            ("rejection_reason_code", "STRING", "否", "查询键", "拒绝原因代码，用于后续频控与分析", "NO_NEED"),
            ("next_follow_up_time", "TIMESTAMP", "否", "查询键", "建议或约定的下次跟进时间", "2026-08-28 10:00:00"),
            ("remark_masked", "STRING", "否", "敏感列", "脱敏跟进摘要；禁止录入完整证件、账号、手机号", "客户希望月底再联系"),
            ("source_system_code", "STRING", "是", "审计", "来源渠道/系统代码", "CRM"),
            ("etl_batch_id", "STRING", "是", "审计", "ETL 批次或流任务检查点标识", "STREAM_88210"),
            ("ingestion_time", "TIMESTAMP", "是", "审计", "进入数仓的时间", "2026-08-25 11:00:08"),
            ("event_dt", "DATE", "是", "分区键", "营销事件分区日期", "2026-08-25"),
        ],
    },
]


VIEWS = [
    ("vw_customer_360_latest", "一名当前有效客户一行", "dim_customer + fct_product_holding + dim_customer_manager", "客户等级、风险/KYC、经理、资产总额、持有产品数、最近快照日", "客户概览与客户经理工作台", "仅暴露脱敏字段；按授权客户集合过滤"),
    ("vw_customer_transaction_summary_12m", "一名客户一行", "fct_transaction", "近 30/90/365 天交易笔数、收入、支出、最近交易日、活跃月数", "交易统计、活跃度筛选", "排除失败交易；默认不显示对手方摘要"),
    ("vw_customer_product_holding_latest", "客户 + 产品一行", "fct_product_holding", "产品类别、余额、市值、收益、到期日、风险等级", "查持仓、到期提醒、交叉销售", "固定最新有效快照；金额按权限控制"),
    ("vw_marketing_campaign_performance", "活动一行", "dim_marketing_campaign + fct_customer_marketing", "目标、触达、响应、转化人数及转化金额、响应率、转化率", "活动效果统计", "采用去重客户数；分母为实际触达客户数"),
    ("vw_customer_marketing_response", "客户 + 活动一行", "fct_customer_marketing", "首次/最近触达、响应、转化、跟进时间、拒绝原因", "查客户参与活动与待跟进事项", "备注只返回脱敏文本"),
    ("vw_manager_customer_portfolio", "客户经理一行", "dim_customer_manager + dim_customer + fct_product_holding", "名下客户数、VIP 数、资产规模、产品覆盖、KYC 到期客户数", "客户经理组合与工作量分析", "经理仅能查看本人/授权团队"),
    ("vw_nl2sql_business_dictionary", "表字段或码值一行", "元数据目录/码表", "中文名、同义词、字段、数据类型、口径、可用聚合、敏感级别", "NL2SQL 术语映射与规则校验", "不包含业务事实数据"),
]


PROCEDURES = [
    ("sp_refresh_customer_snapshot", "p_biz_date DATE, p_run_id STRING, p_force BOOLEAN", "刷新 dim_customer 指定日期快照并维护 SCD2", "校验源数据→去重→关闭旧版本→写入新版本→校验唯一性与行数→发布分区", "每日 02:00", "同一 p_run_id 可安全重试；分区发布前不对外可见"),
    ("sp_load_transaction_partition", "p_biz_date DATE, p_start_ts TIMESTAMP, p_end_ts TIMESTAMP, p_run_id STRING", "装载并对账 fct_transaction 交易分区", "抽取→按 transaction_id 去重→币种折算→冲正关联→写临时分区→源目标对账→交换分区", "小时级 + T+1", "按 transaction_id 幂等；失败保留旧分区"),
    ("sp_refresh_product_holding_snapshot", "p_biz_date DATE, p_run_id STRING", "生成 fct_product_holding 日终持仓快照", "统一产品类别→金额与市值校验→生成 holding_id→写快照→与资产总账核对", "每日 04:00", "仅在对账差异低于阈值时发布"),
    ("sp_merge_customer_marketing_events", "p_event_date DATE, p_checkpoint STRING, p_run_id STRING", "合并触达、响应和转化事件", "读取检查点后事件→按 interaction_id 去重→校验授权与频控标志→写事件分区→更新检查点", "5 分钟/小时", "检查点与分区提交绑定，避免重复/漏数"),
    ("sp_build_campaign_performance", "p_start_date DATE, p_end_date DATE, p_campaign_id STRING, p_run_id STRING", "重算活动效果中间结果或物化视图", "限定日期→去重客户→计算漏斗与转化金额→校验分母→覆盖结果分区", "每日 05:00/按需", "同一活动同一统计窗口全量覆盖"),
    ("sp_collect_nl2sql_stats", "p_biz_date DATE, p_table_name STRING, p_run_id STRING", "收集表/分区统计并生成 NL2SQL 元数据快照", "ANALYZE TABLE→空值/基数/范围统计→敏感字段排除样例→更新业务词典", "装载成功后", "只读扫描；失败不影响事实分区发布"),
]


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("个人金融 NL2SQL 营销查询系统")
    set_run_font(r, size=16, bold=True, color=BLUE)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("数据库设计")
    set_run_font(r2, size=28, bold=True, color=INK)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(36)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Spark SQL + Hive 客户数据仓库")
    set_run_font(r3, size=13, color=DARK_BLUE)

    t = doc.add_table(rows=4, cols=2)
    set_table_geometry(t, [2700, 6660], indent_dxa=120)
    set_table_borders(t, color="D7DBE2")
    cover_rows = [
        ("文档版本", "V1.0（设计初稿）"),
        ("编制日期", "2026-08-25"),
        ("适用范围", "个人金融营销查询、NL2SQL 语义查询、客户经理工作台"),
        ("设计依据", "现有系统整体架构图及仓库项目说明"),
    ]
    for i, (label, value) in enumerate(cover_rows):
        set_cell_shading(t.cell(i, 0), LIGHT_BLUE)
        for j, text in enumerate((label, value)):
            p = t.cell(i, j).paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=(j == 0), color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(48)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = note.add_run("内部使用 · 不含生产客户明文数据")
    set_run_font(rr, size=9.5, color=MUTED)
    doc.add_page_break()


def add_document_body(doc, bullet_num_id):
    add_heading(doc, "1 设计概述", 1)
    add_body(doc, "本设计面向个人金融 NL2SQL 营销查询系统的大数据层。数据通过 Spark SQL 执行，存储于 Hive 兼容客户数据仓库，上层由 Spring Boot 完成身份认证、客户范围校验、敏感信息脱敏、SQL 安全检查、异步执行和结果封装。数据库设计的首要目标是让自然语言查询建立在稳定、可解释、可审计的业务语义上，同时避免模型直接接触明文敏感数据或无界扫描底层明细。")
    add_callout(doc, "核心结论", "六张表采用“客户/客户经理/营销活动维度 + 交易/持仓/客户营销事件事实”的主题模型。Hive 不依赖传统 B-Tree 索引，查询加速通过分区裁剪、分桶、排序、ORC 统计和受控汇总视图实现。")

    add_heading(doc, "1.1 设计范围", 2)
    for text in [
        "设计图中六张核心业务表的中文描述、物理表名、数据粒度、字段、字段说明、主键/关联键、分区与加速策略。",
        "制定适用于 Spark SQL + Hive 的数据库命名、类型、主键、分区、敏感数据、质量、生命周期和变更规范。",
        "设计面向 NL2SQL 和客户经理工作台的数据视图，统一常用指标、关联路径和安全边界。",
        "设计存储过程接口及其 Spark SQL 等价实现，覆盖快照刷新、增量装载、活动效果计算和统计信息收集。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "1.2 边界与待确认假设", 2)
    assumptions = [
        ("交易主题表", "本设计将其解释为交易明细主题事实表，一笔入账交易一行，而不是交易类型维表。"),
        ("产品主数据", "六张表中没有独立产品表，因此 product_id 先关联外部产品主数据；若后续需要完全自包含，应新增 dim_product。"),
        ("机构与权限", "branch_id、org_id 由统一机构主数据解释；行级客户权限仍由接入层强制校验，视图不能替代服务端授权。"),
        ("快照口径", "客户、客户经理、活动使用 SCD2 + 日快照；产品持有使用日终快照；交易和营销事件使用按日增量事实。"),
        ("技术能力", "假设 Spark 3.x 与 Hive Metastore 可用；如不支持 ACID MERGE，则采用临时分区 + 校验 + 原子交换/覆盖分区。"),
        ("数据保留", "建议交易及营销事件在线保留不少于 5 年，最终年限以金融机构监管、审计和内部制度为准。"),
    ]
    add_table(doc, ["假设项", "当前设计取值"], assumptions, [2100, 7260], font_size=9.2)

    add_heading(doc, "2 总体数据架构", 1)
    add_body(doc, "架构图体现了从客户经理工作台、Spring Boot 接入层、NL2SQL 核心业务层到 Spark SQL/Hive 数据仓库的完整链路。数据库层既是查询执行的数据来源，也是权限范围、脱敏字段和业务术语映射的最终落点。")
    if ARCH_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        inline = run.add_picture(str(ARCH_IMAGE), width=Inches(6.45))
        inline._inline.docPr.set("descr", "个人金融 NL2SQL 营销查询系统整体架构图")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        rr = cap.add_run("图 2-1 个人金融 NL2SQL 营销查询系统整体架构")
        set_run_font(rr, size=9, color=MUTED)

    add_heading(doc, "2.1 主题关系", 2)
    relation_rows = [
        ("dim_customer → fct_transaction", "1:N", "customer_id", "一个客户对应多笔交易"),
        ("dim_customer → fct_product_holding", "1:N", "customer_id", "一个客户持有多个账户/产品"),
        ("dim_customer → fct_customer_marketing", "1:N", "customer_id", "一个客户产生多次营销事件"),
        ("dim_marketing_campaign → fct_customer_marketing", "1:N", "campaign_id", "一个活动覆盖多个客户和事件"),
        ("dim_customer_manager → dim_customer", "1:N", "manager_id", "当前客户归属客户经理"),
        ("dim_customer_manager → fct_customer_marketing", "1:N", "manager_id", "客户经理执行触达或跟进"),
        ("外部产品主数据 → 持仓/交易/活动", "1:N", "product_id", "统一解释产品名称、类别和风险等级"),
    ]
    add_table(doc, ["关系", "基数", "关联字段", "说明"], relation_rows, [2700, 900, 1800, 3960], font_size=9)

    add_heading(doc, "2.2 数据分层与访问路径", 2)
    add_body(doc, "六张表作为 NL2SQL 可查询的受控主题层，不直接等同于源系统 ODS。源数据完成标准化、去重、码值统一和敏感字段脱敏后进入主题表；常见复杂关联由稳定视图封装；最终 SQL 由接入层注入日期、客户范围和结果行数限制后交由 Spark SQL 执行。")

    add_heading(doc, "3 数据库设计规范", 1)
    standards = [
        ("库与表命名", "数据库建议命名 pf_nl2sql；维表使用 dim_，明细/快照事实表使用 fct_，视图使用 vw_，过程接口使用 sp_。全部采用小写 snake_case。"),
        ("字段命名", "主标识使用 *_id，代码使用 *_code，金额使用 *_amount，标志使用 *_flag，日期使用 *_date/*_dt，时间使用 *_time，脱敏字段显式带 _masked 或 _hash。"),
        ("数据类型", "标识默认 STRING；计数 BIGINT；金额 DECIMAL(20,2)；利率 DECIMAL(9,6)；日期 DATE；事件时间 TIMESTAMP；禁止用 DOUBLE 存金额。"),
        ("时间口径", "业务时间统一使用 Asia/Hong_Kong；分区日期采用业务日期。跨时区源数据先保留原始时区，再转换后入主题层。"),
        ("主键与唯一性", "Hive 约束不作为强制校验。所有主键/唯一键必须由装载任务执行重复检查并写入质量结果；事实表以业务流水号幂等。"),
        ("空值与默认值", "未知值优先使用 NULL；业务上确有“未知/其他”的代码使用 UNKNOWN/OTHER。禁止用 0、空串或 9999 混淆缺失值。SCD2 当前结束时间例外使用 9999-12-31。"),
        ("码表与口径", "枚举字段必须登记码表、中文名、同义词、有效期和负责人；字段注释中避免同一术语多义。NL2SQL 只允许使用已发布码值。"),
        ("分区设计", "大事实表按业务日期分区；查询网关要求明细查询必须带分区范围。避免按高基数字段分区，也避免创建大量小文件。"),
        ("索引与加速", "Hive 不依赖传统二级索引。优先使用分区裁剪、分桶、桶内排序、ORC/Parquet 统计、Bloom Filter、动态分区裁剪和汇总视图。"),
        ("存储格式", "建议 ORC + ZLIB/SNAPPY；单文件目标 256-512 MB。压缩算法、条带大小和小文件合并策略由平台基准测试确定。"),
        ("敏感数据", "主题层不保存完整姓名、证件号、手机号、账号和交易对手信息；只保留令牌、哈希、后四位或脱敏展示值。敏感列建立分级标签。"),
        ("权限控制", "Spark SQL 账户只读；Spring Boot 按客户经理、机构和授权客户集合注入行级条件。禁止由大模型自由生成权限条件或绕过脱敏视图。"),
        ("审计与血缘", "每表至少包含 source_system_code、etl_batch_id、data_update_time/ingestion_time 和业务日期；记录源目标行数、异常数、代码版本与回退分区。"),
        ("数据质量", "必检项包括主键重复、必填空值、金额范围、时间先后、码值合法、关联完整性、源目标对账、快照波动和小文件数量。"),
        ("生命周期", "热分区、冷分区和归档分层管理；删除或归档必须符合审计和监管要求。测试、日志、提示词和截图不得包含真实客户明文。"),
        ("变更管理", "新增字段向后兼容；重命名、改类型或改口径必须经过版本评审，先更新元数据/视图/契约，再发布任务并提供回退方案。"),
        ("NL2SQL 友好性", "字段注释写明业务名、同义词、单位、分母、时间窗口和禁止组合；复杂指标封装为视图，避免模型重复拼接高风险关联。"),
    ]
    add_table(doc, ["规范类别", "设计要求"], standards, [2100, 7260], font_size=8.8)

    add_heading(doc, "3.1 索引与物理组织原则", 2)
    add_body(doc, "Spark SQL/Hive 的性能优化重点不是传统关系数据库索引，而是减少扫描数据量和 Shuffle。文档中“索引”统一指可执行的物理加速策略：分区等价于一级目录索引，分桶和排序优化关联/聚合，ORC 统计与 Bloom Filter 支持文件和条带跳过，物化汇总减少重复计算。")
    index_rows = [
        ("分区裁剪", "txn_dt/event_dt/snapshot_dt", "所有事实与快照表", "强制日期条件，收益最高"),
        ("分桶", "customer_id 或 campaign_id", "大表等值关联、聚合", "桶数按数据量和集群并发复核"),
        ("桶内排序", "主关联键 + 时间", "最近记录、时间范围查询", "改善局部扫描与窗口函数"),
        ("Bloom Filter", "流水号、客户号、活动号", "高选择性等值查询", "仅对平台支持的 ORC 列启用"),
        ("列统计", "基数、空值、最小/最大值", "优化器选择 Join 顺序", "每次分区发布后收集"),
        ("汇总/物化视图", "客户 12 月交易、活动漏斗", "高频统计", "明确刷新时间和口径版本"),
    ]
    add_table(doc, ["策略", "建议字段", "适用场景", "注意事项"], index_rows, [1600, 2200, 2500, 3060], font_size=9)


def add_core_tables(doc):
    for idx, spec in enumerate(TABLES, start=1):
        sec = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(sec, landscape=True)
        add_heading(doc, f"4.{idx} {spec['cn']}（{spec['name']}）", 2)
        add_body(doc, spec["desc"])
        summary = [
            ("数据粒度", spec["grain"]),
            ("更新方式", spec["refresh"]),
            ("主键/唯一性", spec["key"]),
            ("分区设计", spec["partition"]),
            ("索引/加速", spec["accel"]),
        ]
        add_caption(doc, f"表 4-{idx}-1  {spec['cn']}设计摘要")
        add_table(doc, ["设计项", "设计内容"], summary, [2400, 10560], font_size=8.8, landscape=True)
        add_caption(doc, f"表 4-{idx}-2  {spec['cn']}字段定义")
        rows = []
        for no, field in enumerate(spec["fields"], start=1):
            rows.append((no, *field))
        add_table(
            doc,
            ["序号", "字段名", "数据类型", "必填", "键/索引角色", "字段描述与口径", "示例"],
            rows,
            [550, 2150, 1450, 650, 1500, 4650, 2010],
            font_size=7.8,
            center_cols={0, 2, 3, 4},
            landscape=True,
        )
        add_body(doc, "约束说明：Hive 中 PK、UK、FK 均为逻辑约束，装载任务必须通过数据质量规则保证；除明确列出的脱敏字段外，不得在本主题表增加可还原的客户隐私字段。")


def add_views_and_procedures(doc, bullet_num_id):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(sec, landscape=False)
    add_heading(doc, "5 数据视图设计", 1)
    add_body(doc, "视图用于固化稳定关联、指标口径和脱敏边界，减少 NL2SQL 直接拼接底层事实表的复杂度。视图本身不是权限边界：接入层仍必须校验 JWT、客户经理身份、机构范围和授权客户集合，并对结果执行二次脱敏。")
    view_rows = [(i, *v) for i, v in enumerate(VIEWS, start=1)]
    add_caption(doc, "表 5-1  建议数据视图清单")
    add_table(doc, ["序号", "视图名", "粒度", "来源", "主要输出", "使用场景", "安全/口径"], view_rows,
              [500, 1900, 1100, 1700, 2100, 1200, 860], font_size=7.7, center_cols={0})

    add_heading(doc, "5.1 视图统一规则", 2)
    for text in [
        "最新快照必须显式使用 max(snapshot_dt) 或由调度传入业务日期，不允许使用系统当前日期替代已发布数据日期。",
        "统计人数默认使用 count(distinct customer_id)；交易笔数使用去重后的 transaction_id；金额统一注明原币或人民币。",
        "响应率 = 响应客户数 / 实际触达客户数；转化率 = 转化客户数 / 实际触达客户数。分母为 0 时返回 NULL，不返回 0。",
        "底层视图仅暴露哈希、令牌或脱敏值；手机号、证件号、账号和对手方信息不进入 NL2SQL 业务词典。",
        "所有高数据量视图必须支持日期条件下推；禁止在分区字段上套 date_format、cast 等导致裁剪失效的函数。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "5.2 代表性视图 SQL 骨架", 2)
    add_body(doc, "以下 SQL 为逻辑骨架，部署时应将“最新已发布分区”替换为元数据服务提供的确定日期，并根据实际 Spark/Hive 版本调整语法。")
    add_code_block(doc, """
CREATE OR REPLACE VIEW pf_nl2sql.vw_customer_360_latest AS
WITH latest_holding AS (
  SELECT customer_id,
         SUM(COALESCE(market_value, balance_amount, 0)) AS total_asset_amount,
         COUNT(DISTINCT product_id) AS product_count,
         MAX(snapshot_dt) AS holding_snapshot_dt
  FROM pf_nl2sql.fct_product_holding
  WHERE snapshot_dt = ${latest_holding_dt}
    AND holding_status_code = 'ACTIVE'
  GROUP BY customer_id
)
SELECT c.customer_id, c.customer_name_masked, c.customer_level_code,
       c.risk_level_code, c.kyc_status_code, c.branch_id, c.manager_id,
       m.manager_name_masked,
       COALESCE(h.total_asset_amount, 0) AS total_asset_amount,
       COALESCE(h.product_count, 0) AS product_count,
       c.snapshot_dt AS customer_snapshot_dt, h.holding_snapshot_dt
FROM pf_nl2sql.dim_customer c
LEFT JOIN latest_holding h ON c.customer_id = h.customer_id
LEFT JOIN pf_nl2sql.dim_customer_manager m
  ON c.manager_id = m.manager_id AND m.is_current = true
WHERE c.is_current = true
  AND c.snapshot_dt = ${latest_customer_dt};
""")
    add_code_block(doc, """
CREATE OR REPLACE VIEW pf_nl2sql.vw_marketing_campaign_performance AS
SELECT campaign_id,
       COUNT(DISTINCT CASE WHEN event_type_code IN ('CONTACT','EXPOSURE')
                           AND contact_result_code = 'SUCCESS' THEN customer_id END) AS contacted_customer_count,
       COUNT(DISTINCT CASE WHEN response_flag THEN customer_id END) AS response_customer_count,
       COUNT(DISTINCT CASE WHEN conversion_flag THEN customer_id END) AS conversion_customer_count,
       SUM(CASE WHEN conversion_flag THEN COALESCE(conversion_amount,0) ELSE 0 END) AS conversion_amount,
       COUNT(DISTINCT CASE WHEN response_flag THEN customer_id END)
         / NULLIF(COUNT(DISTINCT CASE WHEN event_type_code IN ('CONTACT','EXPOSURE')
                                     AND contact_result_code='SUCCESS' THEN customer_id END), 0) AS response_rate
FROM pf_nl2sql.fct_customer_marketing
WHERE event_dt BETWEEN ${start_dt} AND ${end_dt}
GROUP BY campaign_id;
""")

    add_heading(doc, "5.3 行列权限落地", 2)
    add_body(doc, "推荐由 Spring Boot 查询网关把授权客户集合写入本次 Spark 会话可见的临时授权表（或使用平台原生 Ranger/湖仓行级策略），再将生成 SQL 包裹为内连接。禁止把 manager_id = 当前用户 作为唯一权限规则，因为代管、团队授权、转岗和临时授权可能导致权限与客户归属不完全相同。")

    add_heading(doc, "6 存储过程设计", 1)
    add_callout(doc, "技术说明", "原生 Hive/Spark SQL 不提供与 Oracle/MySQL 等价的通用 CREATE PROCEDURE 能力。本章定义的是稳定的过程接口，建议由 Spark 作业、调度平台或 Spring Boot 管理服务实现；接口名、参数、幂等规则和发布条件保持“存储过程式”契约。", fill=PALE_GOLD)
    proc_rows = [(i, *p) for i, p in enumerate(PROCEDURES, start=1)]
    add_caption(doc, "表 6-1  数据处理过程接口")
    add_table(doc, ["序号", "过程名", "输入参数", "目的", "主要步骤", "调度", "幂等/发布"], proc_rows,
              [480, 1800, 1900, 1400, 1900, 820, 1060], font_size=7.5, center_cols={0})

    h61 = add_heading(doc, "6.1 通用过程契约", 2)
    # The preceding dense procedure matrix nearly fills a page. An explicit
    # break avoids Word placing this heading inside the header area during
    # pagination on some Windows builds.
    h61.paragraph_format.page_break_before = True
    contract_rows = [
        ("输入", "业务日期/时间窗、run_id、可选 force 标志；所有时间参数必须带明确业务时区。"),
        ("输出", "状态码、影响行数、拒绝行数、质量检查结果、已发布分区、错误摘要和血缘批次。"),
        ("幂等", "同一业务日期 + run_id 重试不得产生重复数据；事实表按业务流水去重，快照表按分区覆盖。"),
        ("原子发布", "先写临时路径/临时分区并校验，通过后再交换或覆盖正式分区；失败时保留上一已发布版本。"),
        ("并发", "同一表同一分区只允许一个活动运行实例；锁由调度平台或元数据运行表维护。"),
        ("错误处理", "数据质量失败与技术失败分开记录；超过阈值时中止发布并告警，不自动吞掉异常。"),
        ("审计", "记录 run_id、代码版本、参数、开始/结束时间、源目标行数、校验结果和回退位置。"),
    ]
    add_table(doc, ["契约项", "要求"], contract_rows, [1800, 7560], font_size=9)

    add_heading(doc, "6.2 过程实现模板", 2)
    add_code_block(doc, """
PROCEDURE sp_load_transaction_partition(p_biz_date, p_start_ts, p_end_ts, p_run_id):
  1. acquire_lock('fct_transaction', p_biz_date, p_run_id)
  2. assert parameters are valid and the source watermark is available
  3. write normalized, deduplicated records to tmp/fct_transaction/txn_dt=p_biz_date/run_id=p_run_id
  4. run checks: transaction_id uniqueness, required fields, amount range,
                 legal code values, source/target reconciliation, file count
  5. if any blocking check fails: mark FAILED; keep current published partition; alert owner
  6. atomically replace or exchange txn_dt=p_biz_date with the validated temporary partition
  7. analyze table partition, record lineage and metrics, mark SUCCESS
  8. release_lock()
""")
    add_body(doc, "如部署平台支持 Iceberg/Delta/Hudi，可将分区交换替换为快照提交和回滚；无论采用哪种表格式，都应保持上述输入、幂等、质量门禁和审计契约一致。")

    add_heading(doc, "7 数据质量与验收", 1)
    quality_rows = [
        ("唯一性", "六张表逻辑主键重复数为 0；SCD2 每个业务键最多一个 is_current=true 版本", "阻断发布"),
        ("完整性", "必填字段空值率为 0；customer_id、campaign_id、manager_id 关联缺失率在批准阈值内", "阻断/告警"),
        ("合理性", "金额精度正确；amount >= 0；活动 start_time < end_time；响应/转化时间不早于触达", "阻断发布"),
        ("一致性", "transaction_date = txn_dt；event_date = event_dt；最新快照日期与发布元数据一致", "阻断发布"),
        ("对账", "交易金额/笔数、持仓余额、营销事件数与源系统对账，差异有原因码", "阈值控制"),
        ("波动", "客户数、资产规模、交易量、触达量较 7/30 日基线异常时告警", "告警/人工确认"),
        ("隐私", "敏感字段扫描不得发现完整手机号、证件号、账号；日志和样例不得包含客户明文", "阻断发布"),
        ("性能", "典型查询命中分区；无无界全表扫描；P95 查询目标与架构中的 20 秒目标一致", "验收"),
    ]
    add_table(doc, ["检查类别", "验收规则", "处理"], quality_rows, [1600, 6100, 1660], font_size=9)

    add_heading(doc, "8 NL2SQL 使用约束", 1)
    for text in [
        "优先让模型查询已发布视图；只有视图无法表达且安全规则允许时，才开放六张底层表。",
        "元数据中为客户、交易、活动、产品持有、客户经理、触达、响应、转化等词登记同义词与唯一字段映射。",
        "生成 SQL 必须通过字段白名单、只读语句、分区范围、最大扫描量、最大返回行数和超时检查。",
        "涉及手机号、证件号、账号、客户姓名或交易对手的自然语言请求，返回脱敏字段或拒绝，不允许模型猜测明文列。",
        "结果解释必须携带数据日期、统计窗口、金额币种和指标分母；用户追问时沿用同一口径。",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "9 后续确认项", 1)
    confirm_rows = [
        ("交易主题粒度", "确认是否确为逐笔交易，而非按客户/日期/类型聚合的交易主题宽表。", "影响字段数量、存储量和响应性能"),
        ("产品主数据", "确认是否已有产品表及产品名称、类别、风险等级、期限等字段。", "决定是否新增 dim_product 及视图关联"),
        ("权限来源", "确认客户经理授权范围来自 CRM、统一权限中心还是 Ranger。", "决定行级权限的实现方式"),
        ("引擎与表格式", "确认 Spark/Hive 版本以及是否使用 Iceberg/Delta/Hudi。", "决定 MERGE、快照回滚和存储过程实现"),
        ("实时性", "确认交易和营销事件的目标延迟，以及 20 秒目标适用的查询数据规模。", "决定是否增加实时汇总或缓存层"),
        ("监管保留期", "确认交易、营销触达和查询审计的保存年限与删除规则。", "决定冷热分层和归档策略"),
    ]
    add_table(doc, ["待确认主题", "建议确认内容", "影响"], confirm_rows, [1800, 4800, 2760], font_size=9)
    add_callout(doc, "版本说明", "本版已可作为数据库设计章节初稿和评审底稿。上述确认项不会影响六张表的核心关系，但会影响最终 DDL、调度参数和性能配置。")


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "个人金融NL2SQL营销查询系统数据库设计"
    doc.core_properties.subject = "Spark SQL + Hive 客户数据仓库设计"
    doc.core_properties.author = "项目组"
    doc.core_properties.keywords = "NL2SQL, Hive, Spark SQL, 数据库设计, 个人金融, 营销查询"
    set_styles(doc)
    for section in doc.sections:
        configure_section(section, landscape=False)
    bullet_num_id = add_numbering_definitions(doc)

    add_cover(doc)
    add_document_body(doc, bullet_num_id)

    add_heading(doc, "4 核心业务表结构设计", 1)
    add_body(doc, "本章六张表均位于建议数据库 pf_nl2sql。为便于字段字典阅读，字段表采用横向页面；“键/索引角色”是逻辑设计标识，不代表 Hive 自动强制主外键。")
    add_core_tables(doc)
    add_views_and_procedures(doc, bullet_num_id)

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
